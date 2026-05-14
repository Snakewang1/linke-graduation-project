"""Insert 17 references into 已修复.docx: sentence-boundary, superscript, no adjacency."""
import docx, re, copy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'D:\cursor项目-毕业设计\LinkE协同_毕业论文_已修复.docx'
doc = docx.Document(SRC)

ref_start = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == '参考文献':
        ref_start = i
        break
print(f"Ref at P{ref_start}, total paras: {len(doc.paragraphs)}")

# Insertion plan: (para_idx, ref_num, sentence_index) -- sentence_index: 0=1st, -1=last
plan = [
    # Ch1 (P132-P153)
    (134, 9, 0), (134, 1, -1),          # 1.1 研究背景
    (138, 11, 0), (138, 12, 1), (138, 13, -1),  # 1.2 研究意义
    (141, 10, -1),                       # 1.3.1 国外现状
    (143, 2, 0), (143, 3, -1),          # 1.3.2 国内现状
    (147, 1, -1),                        # 1.4 主要工作

    # Ch2 (P155-P170)
    (161, 15, 0), (161, 17, -1),        # 2.3 Capacitor
    (163, 14, -1),                       # 2.4 DeepSeek
    (165, 5, 0), (165, 6, -1),          # 2.5 OpenAPI
    (167, 16, -1),                       # 2.6 RBAC
    (169, 5, -1),                        # 2.7 小结

    # Ch3 (P171-P219)
    (173, 9, -1),                        # 3.1 业务需求
    (176, 3, -1),                        # 3.2.1 admin
    (185, 2, -1),                        # 3.3.1 工作台
    (187, 4, -1),                        # 3.3.2 协同
    (191, 14, -1),                       # 3.3.4 AI
    (193, 8, -1),                        # 3.3.5 集成
    (196, 8, -1),                        # 3.4.1 集成流程
    (200, 4, -1),                        # 3.4.2 状态流转
    (218, 2, 0), (218, 3, 1), (218, 4, 2),  # 3.7 小结

    # Ch4 (P220-P265)
    (222, 3, 0), (222, 11, 1), (222, 16, -1),   # 4.1 设计原则
    (224, 5, 0), (224, 7, -1),                    # 4.2 分层架构
    (228, 11, 0), (228, 17, -1),                  # 4.3 响应式
    (232, 3, 0), (232, 7, 1), (232, 10, -1),     # 4.4 模块
    (236, 16, -1),                                  # 4.5 Prompt
    (264, 5, 0), (264, 15, -1),                    # 4.7 小结

    # Ch5 (P266-P336)
    (268, 6, 0), (268, 15, -1),                    # 5.1 项目结构
    (272, 2, -1),                                   # 5.2 仪表盘
    (288, 4, 0), (288, 7, -1),                     # 5.3 协同
    (292, 2, -1),                                   # 5.4 消息
    (296, 14, -1),                                  # 5.5 AI
    (318, 16, -1),                                  # 5.6 权限
    (322, 8, 0), (322, 7, 1), (322, 10, -1),       # 5.7 OpenAPI
    (335, 6, 0), (335, 17, -1),                     # 5.8 小结

    # Ch6 (P337-P358)
    (343, 1, 0), (343, 5, -1),                      # 6.2 测试
    (351, 12, 0), (351, 13, -1),                    # 6.4 性能
    (355, 12, 0), (355, 13, -1),                    # 6.5 改进
    (357, 5, -1),                                    # 6.6 小结
]

# Convert to absolute positions
from collections import defaultdict
by_para = defaultdict(list)
for pi, rn, si in plan:
    para = doc.paragraphs[pi]
    full = para.text
    if len(full) < 40:
        continue
    sent_ends = [m.end() for m in re.finditer(r'[。！？]', full)]
    if not sent_ends:
        by_para[pi].append((len(full), rn))
    elif si == -1:
        by_para[pi].append((sent_ends[-1], rn))
    elif si < len(sent_ends):
        by_para[pi].append((sent_ends[si], rn))
    else:
        by_para[pi].append((sent_ends[-1], rn))

# Deduplicate same (para, pos, ref)
seen = set()
for pi in by_para:
    deduped = []
    for pos, rn in by_para[pi]:
        key = (pi, pos, rn)
        if key not in seen:
            seen.add(key)
            deduped.append((pos, rn))
    by_para[pi] = deduped

print(f"Insertions: {sum(len(v) for v in by_para.values())} across {len(by_para)} paragraphs")

# === PASS 1: Insert [N] text ===
for pi, inserts in by_para.items():
    para = doc.paragraphs[pi]
    for pos, rn in sorted(inserts, key=lambda x: -x[0]):
        p_elem = para._element
        t_info = []
        running = 0
        for r_elem in p_elem.findall(qn('w:r')):
            if r_elem.find(qn('w:fldChar')) is not None:
                continue
            if r_elem.find(qn('w:instrText')) is not None:
                continue
            for t_elem in r_elem.findall(qn('w:t')):
                txt = t_elem.text or ''
                t_info.append((t_elem, running, running + len(txt)))
                running += len(txt)
        if not t_info:
            continue
        target = None
        for ti in t_info:
            if ti[1] <= pos <= ti[2]:
                target = ti
                break
        if target is None:
            if pos >= t_info[-1][2]:
                target = t_info[-1]
                pos = t_info[-1][2]
            else:
                continue
        t_elem, t_start, t_end = target
        offset = max(0, min(pos - t_start, len(t_elem.text or '')))
        old = t_elem.text or ''
        t_elem.text = old[:offset] + f'[{rn}]' + old[offset:]

print("PASS 1: text inserted")

# === PASS 2: Superscript + anti-adjacency via paragraph rebuild ===
for i in range(ref_start):
    para = doc.paragraphs[i]
    full = para.text
    if '[' not in full:
        continue
    p_elem = para._element

    runs_content = []
    for r_elem in p_elem.findall(qn('w:r')):
        rPr = r_elem.find(qn('w:rPr'))
        for t_elem in r_elem.findall(qn('w:t')):
            txt = t_elem.text or ''
            if txt:
                runs_content.append((txt, rPr))
    if not runs_content:
        continue

    new_runs = []
    for txt, rPr in runs_content:
        parts = re.split(r'(\[\d+\])', txt)
        for part in parts:
            if not part:
                continue
            is_cite = re.fullmatch(r'\[\d+\]', part) is not None
            new_runs.append((part, rPr, is_cite))

    # Insert spaces between consecutive citations
    final_runs = []
    for j, (txt, rPr, is_cite) in enumerate(new_runs):
        if j > 0 and is_cite and final_runs[-1][2]:
            final_runs.append((' ', None, False))
        final_runs.append((txt, rPr, is_cite))

    for r_elem in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r_elem)

    for txt, rPr, is_cite in final_runs:
        nr = OxmlElement('w:r')
        if rPr is not None:
            nrPr = copy.deepcopy(rPr)
            if is_cite:
                va = OxmlElement('w:vertAlign')
                va.set(qn('w:val'), 'superscript')
                nrPr.append(va)
            nr.append(nrPr)
        elif is_cite:
            nrPr = OxmlElement('w:rPr')
            va = OxmlElement('w:vertAlign')
            va.set(qn('w:val'), 'superscript')
            nrPr.append(va)
            nr.append(nrPr)
        nt = OxmlElement('w:t')
        nt.set(qn('xml:space'), 'preserve')
        nt.text = txt
        nr.append(nt)
        p_elem.append(nr)

print("PASS 2: superscript applied")

# === SAVE ===
doc.save(SRC)

# === VERIFY ===
doc2 = docx.Document(SRC)
ref_start2 = None
for i, para in enumerate(doc2.paragraphs):
    if para.text.strip() == '参考文献':
        ref_start2 = i
        break

total = 0
adj = 0
ref_usage = {}
for i in range(ref_start2):
    text = doc2.paragraphs[i].text
    cites = re.findall(r'\[(\d+)\]', text)
    total += len(cites)
    adj += len(re.findall(r'\[\d+\]\[\d+\]', text))
    for c in cites:
        ref_usage[int(c)] = ref_usage.get(int(c), 0) + 1

print(f"\nTotal citations: {total}, Adjacent: {adj}")
for n in range(1, 18):
    c = ref_usage.get(n, 0)
    print(f"  [{n}]: {c}x {'OK' if c > 0 else 'MISSING!'}")
print(f"Saved: {SRC}")
